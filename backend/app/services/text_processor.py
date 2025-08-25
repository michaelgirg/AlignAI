import re
import hashlib
import unicodedata
import logging
from typing import List, Tuple, Optional, Dict
import pdfminer.high_level
from docx import Document as DocxDocument
from ..models.models import DocumentSection

logger = logging.getLogger(__name__)

class TextProcessor:
    def __init__(self):
        """Initialize text processor with section detection patterns."""
        self.section_patterns = [
            # Common section headers
            (r'^(?:SUMMARY|PROFILE|OBJECTIVE|ABOUT|INTRODUCTION)', 'summary'),
            (r'^(?:EXPERIENCE|WORK HISTORY|EMPLOYMENT|CAREER|PROFESSIONAL)', 'experience'),
            (r'^(?:EDUCATION|ACADEMIC|QUALIFICATIONS|DEGREES)', 'education'),
            (r'^(?:SKILLS|TECHNOLOGIES|TOOLS|COMPETENCIES|EXPERTISE)', 'skills'),
            (r'^(?:PROJECTS|PORTFOLIO|ACHIEVEMENTS|ACCOMPLISHMENTS)', 'projects'),
            (r'^(?:CERTIFICATIONS|CERTIFICATES|TRAINING)', 'certifications'),
            (r'^(?:LANGUAGES|LANGUAGE SKILLS)', 'languages'),
            (r'^(?:INTERESTS|HOBBIES|ACTIVITIES)', 'interests'),
            (r'^(?:REFERENCES|REFEREES)', 'references'),
            (r'^(?:AWARDS|HONORS|RECOGNITIONS)', 'awards'),
            
            # Alternative patterns with different formatting
            (r'^[•\-\*]\s*(?:SUMMARY|PROFILE|OBJECTIVE)', 'summary'),
            (r'^[•\-\*]\s*(?:EXPERIENCE|WORK HISTORY)', 'experience'),
            (r'^[•\-\*]\s*(?:EDUCATION|ACADEMIC)', 'education'),
            (r'^[•\-\*]\s*(?:SKILLS|TECHNOLOGIES)', 'skills'),
            
            # Numbered sections
            (r'^\d+\.\s*(?:SUMMARY|PROFILE|OBJECTIVE)', 'summary'),
            (r'^\d+\.\s*(?:EXPERIENCE|WORK HISTORY)', 'experience'),
            (r'^\d+\.\s*(?:EDUCATION|ACADEMIC)', 'education'),
            (r'^\d+\.\s*(?:SKILLS|TECHNOLOGIES)', 'skills'),
        ]

    def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extract text from file based on file type."""
        try:
            if file_type.lower() == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                return self._extract_from_docx(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return ""
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            return ""

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using pdfminer."""
        try:
            text = pdfminer.high_level.extract_text(file_path)
            if not text or not text.strip():
                logger.warning(f"PDF extraction returned empty text: {file_path}")
                return ""
            return text
        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {e}")
            return ""

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file using python-docx."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            text = '\n'.join(text_parts)
            if not text or not text.strip():
                logger.warning(f"DOCX extraction returned empty text: {file_path}")
                return ""
            return text
            
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}")
            return ""

    def normalize_text(self, text: str) -> str:
        """Normalize text by cleaning and standardizing format."""
        if not text:
            return ""
        
        try:
            # Unicode normalization
            text = unicodedata.normalize('NFKC', text)
            
            # Replace common special characters
            text = text.replace('\u201c', '"').replace('\u201d', '"')  # Smart quotes
            text = text.replace('\u2018', "'").replace('\u2019', "'")  # Smart apostrophes
            text = text.replace('\u2013', '-').replace('\u2014', '--')  # En/em dashes
            text = text.replace('\u2022', '•')  # Bullet points
            text = text.replace('\u2023', '•')  # Alternative bullet
            text = text.replace('\u25E6', '•')  # White bullet
            
            # Standardize bullet points and dashes
            text = re.sub(r'^[\s]*[•\-\*\+]\s*', '• ', text, flags=re.MULTILINE)
            text = re.sub(r'[\s]*[•\-\*\+]\s*', ' • ', text)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text)  # Multiple spaces to single
            text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple newlines to double
            text = text.strip()
            
            # Fix common formatting issues
            text = re.sub(r'([a-z])\s*-\s*([a-z])', r'\1-\2', text)  # Fix hyphenated words
            text = re.sub(r'([.!?])\s*([A-Z])', r'\1 \2', text)  # Fix sentence spacing
            
            return text
            
        except Exception as e:
            logger.error(f"Text normalization failed: {e}")
            return text

    def detect_sections(self, text: str) -> List[DocumentSection]:
        """Detect document sections using pattern matching."""
        if not text:
            return []
        
        try:
            sections = []
            lines = text.split('\n')
            current_section = None
            current_content = []
            
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                
                # Check if this line is a section header
                section_name = self._identify_section_header(line)
                
                if section_name:
                    # Save previous section if exists
                    if current_section and current_content:
                        sections.append(DocumentSection(
                            name=current_section,
                            text='\n'.join(current_content).strip(),
                            start_line=i,
                            end_line=i - 1
                        ))
                    
                    # Start new section
                    current_section = section_name
                    current_content = []
                else:
                    # Add line to current section content
                    if current_section:
                        current_content.append(line)
                    else:
                        # Content before first section
                        current_content.append(line)
            
            # Add the last section
            if current_section and current_content:
                sections.append(DocumentSection(
                    name=current_section,
                    text='\n'.join(current_content).strip(),
                    start_line=len(lines) - len(current_content),
                    end_line=len(lines) - 1
                ))
            
            # If no sections detected, create a default structure
            if not sections:
                sections = self._create_default_sections(text)
            
            return sections
            
        except Exception as e:
            logger.error(f"Section detection failed: {e}")
            return self._create_default_sections(text)

    def _identify_section_header(self, line: str) -> Optional[str]:
        """Identify if a line is a section header."""
        line_upper = line.upper()
        
        for pattern, section_name in self.section_patterns:
            if re.match(pattern, line_upper):
                return section_name
        
        return None

    def _create_default_sections(self, text: str) -> List[DocumentSection]:
        """Create default sections when pattern detection fails."""
        try:
            lines = text.split('\n')
            total_lines = len(lines)
            
            # Simple heuristic: split into roughly equal parts
            chunk_size = max(1, total_lines // 4)
            
            sections = []
            
            # Summary section (first chunk)
            summary_end = min(chunk_size, total_lines)
            summary_text = '\n'.join(lines[:summary_end])
            sections.append(DocumentSection(
                name="summary",
                text=summary_text,
                start_line=0,
                end_line=summary_end - 1
            ))
            
            # Experience section (second chunk)
            exp_start = summary_end
            exp_end = min(exp_start + chunk_size, total_lines)
            if exp_start < total_lines:
                exp_text = '\n'.join(lines[exp_start:exp_end])
                sections.append(DocumentSection(
                    name="experience",
                    text=exp_text,
                    start_line=exp_start,
                    end_line=exp_end - 1
                ))
            
            # Skills section (third chunk)
            skills_start = exp_end
            skills_end = min(skills_start + chunk_size, total_lines)
            if skills_start < total_lines:
                skills_text = '\n'.join(lines[skills_start:skills_end])
                sections.append(DocumentSection(
                    name="skills",
                    text=skills_text,
                    start_line=skills_start,
                    end_line=skills_end - 1
                ))
            
            # Education section (remaining lines)
            edu_start = skills_end
            if edu_start < total_lines:
                edu_text = '\n'.join(lines[edu_start:])
                sections.append(DocumentSection(
                    name="education",
                    text=edu_text,
                    start_line=edu_start,
                    end_line=total_lines - 1
                ))
            
            return sections
            
        except Exception as e:
            logger.error(f"Failed to create default sections: {e}")
            # Fallback: single section
            return [DocumentSection(
                name="content",
                text=text,
                start_line=0,
                end_line=len(text.split('\n')) - 1
            )]

    def calculate_text_hash(self, text: str) -> str:
        """Calculate SHA-256 hash of text for deduplication."""
        try:
            if not text:
                return ""
            
            # Normalize text before hashing
            normalized = self.normalize_text(text)
            return hashlib.sha256(normalized.encode('utf-8')).hexdigest()
            
        except Exception as e:
            logger.error(f"Failed to calculate text hash: {e}")
            return ""

    def extract_metadata(self, text: str) -> Dict[str, any]:
        """Extract basic metadata from text."""
        try:
            metadata = {
                "word_count": len(text.split()),
                "line_count": len(text.split('\n')),
                "character_count": len(text),
                "has_sections": len(self.detect_sections(text)) > 1
            }
            
            # Detect language (simple heuristic)
            english_chars = len(re.findall(r'[a-zA-Z]', text))
            total_chars = len(re.findall(r'[a-zA-Z0-9]', text))
            
            if total_chars > 0:
                metadata["english_ratio"] = english_chars / total_chars
            else:
                metadata["english_ratio"] = 0.0
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract metadata: {e}")
            return {
                "word_count": 0,
                "line_count": 0,
                "character_count": 0,
                "has_sections": False,
                "english_ratio": 0.0
            }
